"""Redmine RAG 챗봇 웹 서비스"""


"""
- vector db  읽기
- gemini api 연동
- 웹 인터페이스 연동
- multi-turn 대화 지원
- 세션 별 로그인 추가

"""

from flask import Flask, render_template, request, jsonify, session, Response, stream_with_context
from flask_session import Session
import os
import sys
import logging
import re
import threading
import time
import uuid
import json as _json
import base64 as _base64
import requests as _requests
from datetime import timedelta, datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from rag_engine import RedmineRAG
from utils.rag_utils import RAGHelperMixin
from utils.paperbanana_client import generate_diagram
from config import constants as C
from config import patterns as P
from config.diagram_config import (
    PAPERBANANA_BASE_URL,
    PAPERBANANA_TIMEOUT_SECONDS,
    PAPERBANANA_POLL_INTERVAL,
    PAPERBANANA_REFINEMENT_ITERATIONS,
    PAPERBANANA_DEFAULT_DIAGRAM_TYPE,
    PAPERBANANA_MAX_CONCURRENT,
)

# 도식화 동시 요청 제한 (Semaphore)
_visualize_semaphore = threading.Semaphore(PAPERBANANA_MAX_CONCURRENT)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

template_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'templates')
app = Flask(__name__, template_folder=template_dir)
app.secret_key = os.environ.get("SECRET_KEY", "redmine-rag-secret-key-change-in-production")
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=24)
app.config['SESSION_TYPE'] = 'filesystem'
app.config['SESSION_FILE_DIR'] = os.environ.get("SESSION_FILE_DIR", "/vectordb/flask_sessions")
app.config['SESSION_PERMANENT'] = True
app.config['SESSION_USE_SIGNER'] = True
os.makedirs(app.config['SESSION_FILE_DIR'], exist_ok=True)
Session(app)

logger.info("🚀 Redmine RAG 챗봇 초기화 중...")
_REDMINE_API_KEY = os.environ.get("REDMINE_API_KEY", "")
_REDMINE_URL = os.environ.get("REDMINE_URL", "https://your-redmine.example.com")
# 관리자 이름 목록 (환경변수 ADMIN_USERS로 설정, 쉼표 구분)
# 예: ADMIN_USERS=jksjang,admin → {"jksjang", "admin"}
_raw_admin = os.environ.get("ADMIN_USERS", "")
_ADMIN_USERS: set = {u.strip() for u in _raw_admin.split(",") if u.strip()} if _raw_admin else set()
conversation_db_path = os.environ.get("CONVERSATION_DB_PATH")
if not conversation_db_path:
    default_vectordb = os.environ.get("VECTORDB_PATH", "/vectordb/chroma_db_v0.1.2")
    conversation_db_path = os.path.join(os.path.dirname(default_vectordb), "conversation_db")

rag_engine = RedmineRAG(
    vectordb_path=os.environ.get("VECTORDB_PATH", "/vectordb/chroma_db_v0.2.0"),
    collection_name=os.environ.get("COLLECTION_NAME", "redmine_issues_raw_v4"),
    gemini_api_key=os.environ.get("GEMINI_API_KEY"),
    redmine_url=os.environ.get("REDMINE_URL", "https://your-redmine.example.com"),
    use_case=os.environ.get("USE_CASE", "redmine"),
    conversation_db_path=conversation_db_path
)
crf_engine = None
crf_db_path = os.environ.get(
    "CRF_VECTORDB_PATH",
    "/vectordb/crf_data/chroma_db_v0.3.0"
)
crf_collection_name = os.environ.get("CRF_COLLECTION_NAME", "crf_all_cancers_v0.3.0")
try:
    if os.path.exists(crf_db_path):
        crf_engine = RedmineRAG(
            vectordb_path=crf_db_path,
            collection_name=os.environ.get("COLLECTION_NAME", "redmine_issues_raw_v4"),  # 사용되지 않음 (use_case=crf이므로)
            gemini_api_key=os.environ.get("GEMINI_API_KEY"),
            redmine_url=os.environ.get("REDMINE_URL", "https://your-redmine.example.com"),
            use_case="crf",
            conversation_db_path=conversation_db_path,
            crf_collection_name=crf_collection_name  # CRF 컬렉션 명시
        )
        logger.info(f"✅ CRF 엔진 준비 완료! (컬렉션: {crf_collection_name})")
    else:
        logger.warning(f"⚠️ CRF DB 경로 없음: {crf_db_path}")
except Exception as e:
    logger.warning(f"⚠️ CRF 엔진 초기화 실패: {e}")
logger.info("✅ RAG 엔진 준비 완료!")

@app.route('/')
def index():
    return render_template('chat.html')

@app.route('/chat', methods=['POST'])
def chat():
    try:
        data = request.json
        question = data.get('question', '')
        user_name = data.get('user_name', '')
        conversation_id = data.get('conversation_id')

        if not question:
            return jsonify({"error": "질문이 없습니다"}), 400
        if not user_name:
            return jsonify({"error": "사용자 이름이 필요합니다"}), 400
        if not conversation_id:
            return jsonify({"error": "conversation_id가 필요합니다"}), 400

        # RAGHelperMixin의 SESSION_ID_PREFIX 사용 (중복 제거)
        session_id = f"{RAGHelperMixin.SESSION_ID_PREFIX}{user_name}"

        # conversation_id별 세션 키 분리
        hist_key = f'chat_history_{conversation_id}'
        engine_key = f'last_engine_{conversation_id}'
        turn_key = f'turn_index_{conversation_id}'

        if session.get('user_name') != user_name:
            session['user_name'] = user_name
            session['session_id'] = session_id
            # user_name이 바뀌면 is_admin도 함께 재계산 (관리자 비트 오염 방지)
            session['is_admin'] = bool(_ADMIN_USERS) and (user_name in _ADMIN_USERS)
            logger.info(f"🆕 새 세션: {session_id} (사용자: {user_name})")

        chat_history = session.get(hist_key, [])
        turn_index = session.get(turn_key, 0)
        last_engine = session.get(engine_key, None)  # 이전 엔진 가져오기

        logger.info(f"📝 질문: {question} (히스토리: {len(chat_history)}턴, 이전 엔진: {last_engine})")

        # DOCX 내보내기 트리거 — engine.query/save_conversation 이전 early return
        if P.is_docx_export_request(question):
            count = P.get_docx_export_count(question)  # None=전체, N=N개
            answers = rag_engine.get_latest_answers(session_id, conversation_id,
                                                    count=count, answer_kind="rag")
            if not answers:
                return jsonify({"answer": "먼저 질문을 해주세요. 직전 RAG 답변을 DOCX로 내보냅니다.", "sources": []})

            # 여러 답변이면 합치기, diagram 경로 수집
            diagram_paths = []
            if len(answers) == 1:
                merged_q = answers[0]["question"]
                merged_a = answers[0]["answer"]
                target_turn = answers[0]["turn_index"]
                for d in answers[0].get("diagrams", []):
                    if d.get("file_path") and os.path.exists(d["file_path"]):
                        diagram_paths.append(d["file_path"])
            else:
                merged_q = " / ".join(a["question"] for a in answers)
                merged_a = "\n\n---\n\n".join(
                    f"## {a['question']}\n\n{a['answer']}" for a in answers
                )
                target_turn = answers[-1]["turn_index"]
                for a in answers:
                    for d in a.get("diagrams", []):
                        if d.get("file_path") and os.path.exists(d["file_path"]):
                            diagram_paths.append(d["file_path"])

            export_id = str(uuid.uuid4())
            _write_export_state_atomic(export_id, {
                "export_id": export_id,
                "session_id": session_id,
                "conversation_id": conversation_id,
                "turn_index": target_turn,
                "question": merged_q,
                "answer": merged_a,
                "diagram_paths": diagram_paths,
                "status": "pending",
                "created_at": _dt_now_utc(),
                "updated_at": _dt_now_utc(),
                "docx_model": os.environ.get("DOCX_MODEL", "gemini-2.5-pro-preview-05-06"),
                "docx_path": None,
                "error": None,
                "db_save_failed": False,
            })
            label = "전체" if count is None else f"최근 {len(answers)}개"
            logger.info(f"📄 DOCX export 요청({label}): export_id={export_id} turn={target_turn}")
            return jsonify({
                "answer": f"DOCX 문서를 생성합니다. ({label} 답변 포함)",
                "docx_export": True,
                "export_id": export_id,
                "target_turn_index": target_turn,
                "sources": [],
                "turn_index": turn_index,
            })

        # 1단계: 명시적 키워드 체크 (빠른 경로)
        is_crf_query = crf_engine is not None and crf_engine.is_crf_data_query(question)
        is_redmine_query = bool(rag_engine._extract_issue_ids(question)) or rag_engine.is_redmine_data_query(question)

        # 후속 질문 패턴 감지 (애매한 질문)
        is_followup = any(re.search(p, question, re.IGNORECASE) for p in P.FOLLOWUP_PATTERNS)

        # 2단계: 라우팅 결정
        route_reason = None
        if is_redmine_query:
            # Redmine 이슈 번호나 명시적 키워드가 있으면 Redmine 확정
            engine = rag_engine
            engine_name = 'redmine'
            route_reason = 'explicit_keyword'
            logger.info("🧭 라우팅: Redmine DB (명시적 키워드)")
        elif is_crf_query and not is_redmine_query:
            # CRF 키워드만 있고 Redmine 키워드 없으면 CRF 확정
            engine = crf_engine
            engine_name = 'crf'
            route_reason = 'explicit_keyword'
            logger.info("🧭 라우팅: CRF DB (명시적 키워드)")
        elif last_engine and is_followup and crf_engine is not None:
            # 이전 엔진이 있고 후속 질문이면 이전 엔진 우선 사용
            if last_engine == 'crf':
                engine = crf_engine
                engine_name = 'crf'
                route_reason = 'followup_context'
                logger.info("🧭 라우팅: CRF DB (이전 맥락 유지 - 후속 질문)")
            else:
                engine = rag_engine
                engine_name = 'redmine'
                route_reason = 'followup_context'
                logger.info("🧭 라우팅: Redmine DB (이전 맥락 유지 - 후속 질문)")
        elif crf_engine is not None:
            # 애매한 경우: 벡터 유사도 비교
            logger.info("🧭 애매한 질문 → 벡터 유사도 비교 시작")

            crf_result = crf_engine.compare_collection_similarity(question)
            redmine_result = rag_engine.compare_collection_similarity(question)

            crf_distance = crf_result['distance']
            redmine_distance = redmine_result['distance']

            logger.info(f"  📊 CRF 거리: {crf_distance:.4f} vs Redmine 거리: {redmine_distance:.4f}")

            # 유사도 차이 임계값 (0.05 = 5% 차이)
            threshold = C.ROUTING_SIMILARITY_THRESHOLD

            if crf_distance < redmine_distance - threshold:
                engine = crf_engine
                engine_name = 'crf'
                route_reason = 'vector_similarity'
                logger.info(f"  ✅ CRF DB 선택 (거리 차이: {redmine_distance - crf_distance:.4f})")
            elif redmine_distance < crf_distance - threshold:
                engine = rag_engine
                engine_name = 'redmine'
                route_reason = 'vector_similarity'
                logger.info(f"  ✅ Redmine DB 선택 (거리 차이: {crf_distance - redmine_distance:.4f})")
            else:
                # 거리가 비슷하면 이전 엔진 우선, 없으면 Redmine 기본
                if last_engine == 'crf':
                    engine = crf_engine
                    engine_name = 'crf'
                    route_reason = 'similarity_tie_context'
                    logger.info(f"  ⚖️ 유사도 비슷함 → CRF DB (이전 맥락 유지)")
                else:
                    engine = rag_engine
                    engine_name = 'redmine'
                    route_reason = 'similarity_tie_default'
                    logger.info(f"  ⚖️ 유사도 비슷함 → Redmine DB (default)")
        else:
            engine = rag_engine
            engine_name = 'redmine'
            route_reason = 'no_crf_engine'
            logger.info("🧭 라우팅: Redmine DB (CRF 엔진 없음)")

        result = engine.query(
            question,
            chat_history=chat_history,
            session_id=session_id,
            engine_name=engine_name,
            route_reason=route_reason,
            conversation_id=conversation_id
        )

        sources_summary = result.get('sources', [])
        engine.save_conversation(
            session_id=session_id,
            turn_index=turn_index,
            question=question,
            answer=result['answer'],
            conversation_id=conversation_id,
            sources_summary=sources_summary,
            answer_kind="rag",
        )

        chat_history.append({"question": question, "answer": result['answer']})
        session[hist_key] = chat_history[-C.CHAT_HISTORY_CONFIG["max_turns_in_memory"]:]
        session[turn_key] = turn_index + 1
        session[engine_key] = engine_name  # 현재 사용한 엔진 저장
        session.modified = True

        logger.info(f"✅ 답변 완료 (메모리: {len(chat_history)}턴, 사용 엔진: {engine_name})")
        result['turn_index'] = turn_index
        return jsonify(result)

    except Exception as e:
        logger.error(f"❌ 오류: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/reset', methods=['POST'])
def reset():
    data = request.get_json(silent=True) or {}
    conversation_id = data.get('conversation_id')
    if conversation_id:
        for key in [f'chat_history_{conversation_id}', f'last_engine_{conversation_id}', f'turn_index_{conversation_id}']:
            session.pop(key, None)
        session.modified = True
        logger.info(f"🔄 대화 히스토리 초기화 (conv: {conversation_id})")
    else:
        # 전체 초기화: 대화 히스토리 + 사용자 세션 상태 모두 클리어
        keys_to_remove = [k for k in session.keys() if k.startswith(('chat_history_', 'last_engine_', 'turn_index_'))]
        for k in keys_to_remove:
            session.pop(k, None)
        session.pop('user_name', None)
        session.pop('session_id', None)
        session.pop('is_admin', None)
        session.modified = True
        logger.info("🔄 전체 대화 히스토리 초기화")
    return jsonify({"message": "대화 히스토리가 초기화되었습니다"})

@app.route('/health')
def health():
    return jsonify({
        "status": "healthy",
        "vectordb_count": rag_engine.get_document_count()
    })

@app.route('/session/init', methods=['POST'])
def session_init():
    """프론트 로그인 시 서버 세션에 user_name을 즉시 등록 (이미지 인증용)"""
    data = request.get_json(silent=True) or {}
    user_name = data.get('user_name', '')
    if not user_name:
        return jsonify({"error": "user_name 필요"}), 400
    is_admin = bool(_ADMIN_USERS) and (user_name in _ADMIN_USERS)
    session['user_name'] = user_name
    session['session_id'] = f"{RAGHelperMixin.SESSION_ID_PREFIX}{user_name}"
    session['is_admin'] = is_admin
    session.modified = True
    logger.info(f"🔐 세션 초기화: {user_name} (admin={is_admin})")
    return jsonify({"ok": True, "is_admin": is_admin})


@app.route('/conversations/new', methods=['POST'])
def new_conversation():
    """새 대화 ID 발급"""
    conversation_id = f"conv_{uuid.uuid4().hex[:16]}"
    logger.info(f"🆕 새 대화 생성: {conversation_id}")
    return jsonify({"conversation_id": conversation_id})

@app.route('/conversations', methods=['GET'])
def get_conversations():
    """사용자의 대화 스레드 목록 조회"""
    try:
        user_name = request.args.get('user_name', '')
        if not user_name:
            return jsonify({"error": "user_name이 필요합니다"}), 400
        session_id = f"{RAGHelperMixin.SESSION_ID_PREFIX}{user_name}"
        conversations = rag_engine.get_conversations_list(session_id)
        return jsonify({"conversations": conversations})
    except Exception as e:
        logger.error(f"❌ 대화 목록 조회 오류: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/conversations/<conversation_id>', methods=['GET'])
def get_conversation(conversation_id):
    """특정 대화 스레드 내용 조회"""
    try:
        user_name = request.args.get('user_name', '')
        if not user_name:
            return jsonify({"error": "user_name이 필요합니다"}), 400
        session_id = f"{RAGHelperMixin.SESSION_ID_PREFIX}{user_name}"
        result = rag_engine.get_conversation_by_id(session_id, conversation_id)
        return jsonify(result)
    except Exception as e:
        logger.error(f"❌ 대화 조회 오류: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/conversations/<conversation_id>', methods=['DELETE'])
def delete_conversation(conversation_id):
    """특정 대화 스레드 삭제"""
    try:
        user_name = request.args.get('user_name', '')
        if not user_name:
            return jsonify({"error": "user_name이 필요합니다"}), 400
        session_id = f"{RAGHelperMixin.SESSION_ID_PREFIX}{user_name}"
        # Flask 세션에서 해당 conversation 관련 키 제거
        for prefix in ('chat_history_', 'last_engine_', 'turn_index_'):
            session.pop(f'{prefix}{conversation_id}', None)
        session.modified = True
        success = rag_engine.delete_conversation(session_id, conversation_id)
        if success:
            logger.info(f"🗑️ 대화 삭제: {conversation_id} (user: {user_name})")
            return jsonify({"message": "대화가 삭제되었습니다"})
        return jsonify({"error": "삭제 실패"}), 500
    except Exception as e:
        logger.error(f"❌ 대화 삭제 오류: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/save-diagram', methods=['POST'])
def save_diagram():
    """도식화 이미지를 해당 대화 턴에 저장"""
    try:
        data = request.get_json()
        user_name = data.get('user_name', '')
        conversation_id = data.get('conversation_id', '')
        turn_index = data.get('turn_index')
        image_base64 = data.get('image_base64', '')
        task_id = data.get('task_id', '')
        mode = data.get('mode', 'default')
        if not all([user_name, conversation_id, image_base64, turn_index is not None]):
            return jsonify({"error": "필수 파라미터 누락"}), 400
        # turn_index를 int로 명시 변환
        try:
            turn_index = int(turn_index)
        except (TypeError, ValueError):
            return jsonify({"error": "turn_index must be integer"}), 400
        session_id = f"{RAGHelperMixin.SESSION_ID_PREFIX}{user_name}"
        success = rag_engine.save_diagram_to_turn(
            session_id=session_id,
            conversation_id=conversation_id,
            turn_index=turn_index,
            image_base64=image_base64,
            task_id=task_id,
            mode=mode
        )
        if success:
            return jsonify({"ok": True})
        return jsonify({"error": "저장 실패"}), 500
    except Exception as e:
        logger.error(f"❌ 도식화 저장 오류: {str(e)}")
        return jsonify({"error": str(e)}), 500


@app.route('/diagram-image')
def diagram_image():
    """파일시스템에 저장된 도식화 이미지를 base64로 반환"""
    if not session.get("user_name"):
        return jsonify({"error": "unauthorized"}), 401
    file_path = request.args.get('path', '')
    if not file_path:
        return jsonify({"error": "path 파라미터 필요"}), 400
    # 경로 조작 방지: diagrams 디렉토리 내 파일만 허용
    diagrams_dir = os.environ.get("DIAGRAMS_DIR", "/vectordb/diagrams")
    real_path = os.path.realpath(file_path)
    real_dir = os.path.realpath(diagrams_dir)
    if not real_path.startswith(real_dir + os.sep):
        return jsonify({"error": "invalid path"}), 403
    if not os.path.isfile(real_path):
        return jsonify({"error": "not found"}), 404
    try:
        with open(real_path, "rb") as f:
            b64 = _base64.b64encode(f.read()).decode("utf-8")
        return jsonify({"image_base64": f"data:image/png;base64,{b64}"})
    except Exception as e:
        logger.error(f"❌ 도식화 이미지 읽기 오류: {e}")
        return jsonify({"error": str(e)}), 500

def _require_admin():
    """관리자 세션 검증 — 실패 시 (response, status_code) 튜플 반환, 성공 시 None"""
    if not session.get('is_admin'):
        return jsonify({"error": "관리자 권한이 필요합니다"}), 403
    return None

@app.route('/users', methods=['GET'])
def get_users():
    """사용자 목록 조회 (관리자 전용)"""
    err = _require_admin()
    if err:
        return err
    try:
        users = rag_engine.get_user_list()
        return jsonify({"users": users})
    except Exception as e:
        logger.error(f"❌ 사용자 목록 조회 오류: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/users/<user_name>', methods=['DELETE'])
def delete_user(user_name):
    """사용자 삭제 (관리자 전용)"""
    err = _require_admin()
    if err:
        return err
    try:
        success = rag_engine.delete_user(user_name)
        if success:
            return jsonify({"message": f"사용자 '{user_name}' 삭제 완료"})
        else:
            return jsonify({"error": "삭제 실패"}), 500
    except Exception as e:
        logger.error(f"❌ 사용자 삭제 오류: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/visualize', methods=['POST'])
def visualize():
    """
    1단계: LLM 재작성만 수행 → rewritten 데이터 반환
    프론트가 취소 여부 확인 후 /pb-start 호출

    Request JSON:
        {"question": str, "rag_answer": str, "user_name": str, "mode": str}
        mode: "patent" → 특허 청구항 계층 구조 재작성, 생략/기타 → 일반 도식화
    Response JSON:
        {"source_context": str, "communicative_intent": str}  or  {"error": str}
    """
    request_id = str(uuid.uuid4())[:8]
    t_start = time.time()

    try:
        data = request.json or {}
        question = data.get("question", "").strip()
        rag_answer = data.get("rag_answer", "").strip()
        user_name = data.get("user_name", "unknown")
        mode = data.get("mode", "default")

        if not question or not rag_answer:
            return jsonify({"error": "question과 rag_answer가 필요합니다"}), 400
        if len(question) > 500:
            return jsonify({"error": "question이 너무 깁니다 (최대 500자)"}), 400
        if len(rag_answer) > 5000:
            return jsonify({"error": "rag_answer가 너무 깁니다 (최대 5000자)"}), 400

        logger.info(f"[{request_id}] /visualize 요청: user={user_name}, question={question[:60]}...")

        acquired = _visualize_semaphore.acquire(blocking=False)
        if not acquired:
            logger.warning(f"[{request_id}] 도식화 동시 요청 초과")
            return jsonify({"error": "도식화 요청이 많습니다. 잠시 후 다시 시도하세요."}), 429

        try:
            t1 = time.time()
            rewritten = rag_engine._rewrite_for_diagram(question, rag_answer, mode=mode)
            if not rewritten:
                return jsonify({"error": "도식화 입력 변환에 실패했습니다"}), 500

            source_context = rewritten.get("source_context", "")
            communicative_intent = rewritten.get("communicative_intent", "")
            logger.info(f"[{request_id}] LLM 재작성 완료 ({time.time() - t1:.1f}s)")
        finally:
            _visualize_semaphore.release()

        elapsed = round(time.time() - t_start, 1)
        logger.info(f"[{request_id}] 재작성 완료 ({elapsed}s)")
        return jsonify({
            "source_context": source_context,
            "communicative_intent": communicative_intent,
            "request_id": request_id,
        })

    except Exception as e:
        logger.error(f"[{request_id}] /visualize 예외: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500


@app.route('/pb-start', methods=['POST'])
def pb_start():
    """
    2단계: paperbanana task 시작 → task_id 반환
    프론트가 취소하지 않은 경우에만 호출

    Request JSON:
        {"source_context": str, "communicative_intent": str}
    Response JSON:
        {"task_id": str}  or  {"error": str}
    """
    request_id = str(uuid.uuid4())[:8]
    try:
        data = request.json or {}
        source_context = data.get("source_context", "")
        communicative_intent = data.get("communicative_intent", "")

        if not source_context or not communicative_intent:
            return jsonify({"error": "source_context와 communicative_intent가 필요합니다"}), 400

        payload = {
            "source_context": source_context,
            "communicative_intent": communicative_intent,
            "diagram_type": PAPERBANANA_DEFAULT_DIAGRAM_TYPE,
            "refinement_iterations": PAPERBANANA_REFINEMENT_ITERATIONS,
        }
        resp = _requests.post(
            f"{PAPERBANANA_BASE_URL}/api/generate",
            json=payload,
            timeout=15
        )
        resp.raise_for_status()
        task_id = resp.json().get("task_id")
        if not task_id:
            return jsonify({"error": "paperbanana task_id 없음"}), 502

        logger.info(f"[{request_id}] paperbanana task 시작: task_id={task_id}")
        return jsonify({"task_id": task_id})

    except _requests.exceptions.RequestException as e:
        return _pb_request_error(f"[{request_id}] paperbanana", e)
    except Exception as e:
        logger.error(f"[{request_id}] /pb-start 예외: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500


@app.route('/pb-stream/<task_id>')
def pb_stream(task_id):
    """
    paperbanana SSE를 프록시하여 프론트에 실시간 진행 상태 전달
    완료 이벤트에 elapsed_seconds 필드 추가
    """
    t_start = time.time()

    def generate():
        stream_url = f"{PAPERBANANA_BASE_URL}/api/stream/{task_id}"
        try:
            with _requests.get(stream_url, stream=True, timeout=(10, None)) as sse_resp:
                sse_resp.raise_for_status()
                for raw_line in sse_resp.iter_lines(decode_unicode=True):
                    if not raw_line:
                        yield "\n"
                        continue
                    if raw_line.startswith(":"):
                        yield f"{raw_line}\n\n"
                        continue
                    if raw_line.startswith("data:"):
                        data_str = raw_line[5:].strip()
                        try:
                            event = _json.loads(data_str)
                        except _json.JSONDecodeError:
                            yield f"{raw_line}\n\n"
                            continue

                        # 완료 이벤트에 elapsed_seconds 추가
                        if event.get("done") or event.get("agent") == "Complete":
                            event["elapsed_seconds"] = round(time.time() - t_start, 1)

                        yield f"data: {_json.dumps(event, ensure_ascii=False)}\n\n"

                        if event.get("done") or event.get("agent") == "Complete":
                            break
        except Exception as e:
            err = _json.dumps({"error": str(e), "done": True})
            yield f"data: {err}\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",  # Nginx 버퍼링 비활성화
        }
    )


def _pb_request_error(label: str, e: Exception):
    """pb-* 엔드포인트 공통 RequestException → 502 응답 헬퍼"""
    logger.error(f"{label} 실패: {e}")
    return jsonify({"error": str(e)}), 502


@app.route('/pb-cancel/<task_id>', methods=['DELETE'])
def pb_cancel(task_id):
    """paperbanana 생성 작업 취소 요청을 프록시"""
    try:
        resp = _requests.delete(
            f"{PAPERBANANA_BASE_URL}/api/tasks/{task_id}",
            timeout=10
        )
        return jsonify(resp.json()), resp.status_code
    except _requests.exceptions.RequestException as e:
        return _pb_request_error(f"pb-cancel {task_id}", e)


@app.route('/pb-image/<task_id>')
def pb_image(task_id):
    """
    paperbanana 완성 이미지를 base64로 프록시
    Response JSON: {"image_base64": "data:image/png;base64,..."}
    """
    try:
        img_resp = _requests.get(
            f"{PAPERBANANA_BASE_URL}/api/images/{task_id}",
            timeout=30
        )
        img_resp.raise_for_status()
        b64 = _base64.b64encode(img_resp.content).decode("utf-8")
        return jsonify({"image_base64": f"data:image/png;base64,{b64}"})
    except _requests.exceptions.RequestException as e:
        return _pb_request_error(f"pb-image {task_id}", e)


@app.route('/redmine-image/<int:attachment_id>')
def redmine_image_proxy(attachment_id):
    """Redmine 첨부파일 프록시 — 이미지 인라인 표시 / 그 외 파일 다운로드"""
    if not session.get("user_name"):
        return "", 401
    try:
        url = f"{_REDMINE_URL}/attachments/download/{attachment_id}"
        resp = _requests.get(
            url,
            headers={"X-Redmine-API-Key": _REDMINE_API_KEY},
            timeout=30,
            verify=False,
        )
        if resp.status_code != 200:
            return "", resp.status_code
        content_type = resp.headers.get("Content-Type", "application/octet-stream")
        filename = resp.headers.get("Content-Disposition", "")
        # Content-Disposition 헤더에서 파일명 추출, 없으면 attachment_id 사용
        import re as _re
        fname_match = _re.search(r'filename[^;=\n]*=(["\']?)([^"\'\n;]+)\1', filename)
        fname = fname_match.group(2).strip() if fname_match else f"attachment_{attachment_id}"

        headers = {
            "Content-Type": content_type,
            "Cache-Control": "private, max-age=3600",  # 반복 요청 재다운로드 방지
        }
        if content_type.startswith("image/"):
            # 이미지: 인라인 표시
            headers["Content-Disposition"] = f'inline; filename="{fname}"'
        else:
            # 그 외(PDF, PPT, DOCX 등): 다운로드
            headers["Content-Disposition"] = f'attachment; filename="{fname}"'

        return Response(
            stream_with_context(iter([resp.content])),
            headers=headers
        )
    except Exception as e:
        logger.warning(f"첨부파일 프록시 오류 (id={attachment_id}): {e}")
        return "", 502


# ============================================================
# DOCX 내보내기 — 헬퍼 + 엔드포인트
# ============================================================

import tempfile as _tempfile
from datetime import timezone as _tz  # noqa: F811

_EXPORTS_DIR = os.environ.get("EXPORTS_DIR", "/vectordb/exports")
_DOCX_TASKS: dict = {}  # export_id → threading.Event (SSE polling 알림용)

_DOCX_REWRITE_PROMPT = """당신은 문서 편집 전문가입니다. 아래 RAG 답변을 DOCX 문서용 구조화 마크다운으로 재작성해주세요.

사용자 질문: {question}

RAG 답변:
{rag_answer}

출력 규칙:
- 허용 마크다운: ## / ### 헤더, **굵게**, | 표 | (헤더+---+데이터, 셀 안 | 금지), - 불릿, 일반 텍스트
- code fence 금지, HTML 태그 금지, 번호 리스트 금지
- 수치(AUC/Dice/날짜/버전/%) 원문 그대로 — 변경·추정·반올림 절대 금지
- Issue #번호 출처 유지
- "검색된 문서에서..." 등 메타 발화 제거
- 핵심만, 중복·부연 제거
- 마크다운만 출력, 추가 설명 없음"""


def _dt_now_utc() -> str:
    return datetime.now(_tz.utc).isoformat()


def _is_valid_uuid(val: str) -> bool:
    import re as _re
    return bool(_re.match(
        r'^[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}$',
        str(val), _re.IGNORECASE
    ))


def _write_export_state_atomic(export_id: str, state: dict):
    os.makedirs(_EXPORTS_DIR, exist_ok=True)
    final = os.path.join(_EXPORTS_DIR, f"{export_id}.json")
    fd, tmp = _tempfile.mkstemp(dir=_EXPORTS_DIR, suffix=".tmp")
    try:
        with os.fdopen(fd, "w") as f:
            _json.dump(state, f, ensure_ascii=False)
        os.replace(tmp, final)
    except Exception:
        try:
            os.unlink(tmp)
        except Exception:
            pass
        raise


def _read_export_state(export_id: str) -> dict:
    path = os.path.join(_EXPORTS_DIR, f"{export_id}.json")
    try:
        with open(path) as f:
            return _json.load(f)
    except Exception:
        return {}


def _validate_rewrite(original: str, rewritten: str) -> bool:
    import re as _re
    num_pattern = r'v?[\d,]+\.?\d*%?'
    orig_nums = set(_re.findall(num_pattern, original))
    new_nums = set(_re.findall(num_pattern, rewritten))
    orig_issues = set(_re.findall(r'#\d+', original))
    new_issues = set(_re.findall(r'#\d+', rewritten))
    orig_cols = [len(r.split('|')) for r in original.splitlines() if '|' in r]
    new_cols = [len(r.split('|')) for r in rewritten.splitlines() if '|' in r]
    if orig_nums - new_nums or orig_issues - new_issues or orig_cols != new_cols:
        return False
    return True


def _sanitize_for_docx(text: str) -> str:
    import re as _re
    text = _re.sub(r'```[^\n]*\n.*?```', '', text, flags=_re.DOTALL)
    text = _re.sub(r'<[^>]+>', '', text)
    return text.strip()


def _markdown_to_docx(md_text: str, diagram_paths: list = None):
    from docx import Document
    from docx.shared import Pt
    import re as _re

    doc = Document()

    def _add_bold_run(para, text):
        parts = _re.split(r'\*\*(.+?)\*\*', text)
        for i, part in enumerate(parts):
            run = para.add_run(part)
            if i % 2 == 1:
                run.bold = True

    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # 표 감지
        if '|' in line and i + 1 < len(lines) and re.match(r'^\s*\|?[-| :]+\|', lines[i + 1]):
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                if not re.match(r'^\s*\|?[-| :]+\|', lines[i]):
                    table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 1:
                headers = [c.strip() for c in table_lines[0].strip().strip('|').split('|')]
                rows = []
                for tl in table_lines[1:]:
                    rows.append([c.strip() for c in tl.strip().strip('|').split('|')])
                ncols = len(headers)
                table = doc.add_table(rows=1 + len(rows), cols=ncols)
                table.style = 'Table Grid'
                for j, h in enumerate(headers):
                    cell = table.cell(0, j)
                    cell.text = h
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                for ri, row in enumerate(rows):
                    for j in range(ncols):
                        table.cell(ri + 1, j).text = row[j] if j < len(row) else ''
            continue

        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('- ') or line.startswith('* '):
            para = doc.add_paragraph(style='List Bullet')
            _add_bold_run(para, line[2:].strip())
        elif line.strip():
            para = doc.add_paragraph()
            _add_bold_run(para, line.strip())

        i += 1

    # 도식화 이미지 삽입 (답변 내용 뒤에)
    if diagram_paths:
        from docx.shared import Inches
        doc.add_heading('도식화', level=2)
        for img_path in diagram_paths:
            if os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(5.5))
                except Exception:
                    pass

    return doc


def _run_docx_export(export_id: str):
    lock_path = os.path.join(_EXPORTS_DIR, f"{export_id}.lock")
    try:
        state = _read_export_state(export_id)
        if not state:
            return

        # pending → running
        state["status"] = "running"
        state["updated_at"] = _dt_now_utc()
        _write_export_state_atomic(export_id, state)

        answer = state.get("answer", "")

        # LLM 재작성 없이 원본 답변 그대로 변환
        structured_md = _sanitize_for_docx(answer)
        logger.info(f"  📄 DOCX 변환 시작: {len(structured_md)}자 (export_id={export_id})")

        # running → generating_docx
        state["status"] = "generating_docx"
        state["updated_at"] = _dt_now_utc()
        _write_export_state_atomic(export_id, state)

        diagram_paths = state.get("diagram_paths") or []
        doc = _markdown_to_docx(structured_md, diagram_paths=diagram_paths)
        docx_path = os.path.join(_EXPORTS_DIR, f"{export_id}.docx")
        doc.save(docx_path)
        if diagram_paths:
            logger.info(f"  🖼 도식화 {len(diagram_paths)}개 포함 (export_id={export_id})")

        # generating_docx → complete
        state["status"] = "complete"
        state["docx_path"] = docx_path
        state["updated_at"] = _dt_now_utc()
        _write_export_state_atomic(export_id, state)

        # DB turn에 export 정보 append
        try:
            rag_engine.save_export_to_turn(
                session_id=state["session_id"],
                conversation_id=state["conversation_id"],
                turn_index=state["turn_index"],
                export_id=export_id,
                file_path=docx_path,
            )
        except Exception as db_err:
            logger.warning(f"  ⚠️ export DB 저장 실패 (파일은 완료): {db_err}")
            state["db_save_failed"] = True
            _write_export_state_atomic(export_id, state)

        logger.info(f"  ✅ DOCX export 완료: {docx_path}")

    except Exception as e:
        logger.error(f"❌ DOCX export 실패 (export_id={export_id}): {e}", exc_info=True)
        try:
            state = _read_export_state(export_id) or {}
            state["status"] = "error"
            state["error"] = str(e)
            state["updated_at"] = _dt_now_utc()
            _write_export_state_atomic(export_id, state)
        except Exception:
            pass
    finally:
        try:
            if os.path.exists(lock_path):
                os.unlink(lock_path)
        except Exception:
            pass
        # SSE polling 알림
        ev = _DOCX_TASKS.get(export_id)
        if ev:
            ev.set()


@app.route('/docx-direct', methods=['POST'])
def docx_direct():
    """버튼 클릭으로 특정 turn을 바로 DOCX export (상태파일 생성 후 start까지 원스텝)"""
    try:
        data = request.json or {}
        user_name = data.get("user_name", "")
        conversation_id = data.get("conversation_id", "")
        turn_index = data.get("turn_index")
        answer = data.get("answer", "")
        question = data.get("question", "")

        if not answer:
            return jsonify({"error": "answer 없음"}), 400

        session_id = f"{RAGHelperMixin.SESSION_ID_PREFIX}{user_name}"

        # 해당 turn의 도식화 이미지 경로 수집
        diagram_paths = []
        try:
            turn_data = rag_engine.get_latest_answers(session_id, conversation_id,
                                                      count=None, answer_kind="rag")
            for a in turn_data:
                if a.get("turn_index") == turn_index:
                    for d in a.get("diagrams", []):
                        if d.get("file_path") and os.path.exists(d["file_path"]):
                            diagram_paths.append(d["file_path"])
                    break
        except Exception:
            pass

        export_id = str(uuid.uuid4())
        _write_export_state_atomic(export_id, {
            "export_id": export_id,
            "session_id": session_id,
            "conversation_id": conversation_id,
            "turn_index": turn_index,
            "question": question,
            "answer": answer,
            "diagram_paths": diagram_paths,
            "status": "pending",
            "created_at": _dt_now_utc(),
            "updated_at": _dt_now_utc(),
            "docx_model": os.environ.get("DOCX_MODEL", "gemini-2.5-pro-preview-05-06"),
            "docx_path": None,
            "error": None,
            "db_save_failed": False,
        })

        # 바로 스레드 시작 (lock 획득 포함)
        os.makedirs(_EXPORTS_DIR, exist_ok=True)
        lock_path = os.path.join(_EXPORTS_DIR, f"{export_id}.lock")
        fd = os.open(lock_path, os.O_CREAT | os.O_EXCL | os.O_WRONLY)
        os.write(fd, str(os.getpid()).encode())
        os.close(fd)

        ev = threading.Event()
        _DOCX_TASKS[export_id] = ev
        threading.Thread(target=_run_docx_export, args=(export_id,), daemon=True).start()
        logger.info(f"📄 DOCX direct export 시작: export_id={export_id} turn={turn_index}")
        return jsonify({"export_id": export_id})
    except Exception as e:
        logger.error(f"❌ /docx-direct 오류: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500


@app.route('/docx-start', methods=['POST'])
def docx_start():
    """DOCX export 백그라운드 작업 시작"""
    try:
        data = request.json or {}
        export_id = data.get("export_id", "")
        user_name = data.get("user_name", "")
        if not _is_valid_uuid(export_id):
            return jsonify({"error": "invalid export_id"}), 400

        session_id = f"{RAGHelperMixin.SESSION_ID_PREFIX}{user_name}"
        state = _read_export_state(export_id)
        if not state:
            return jsonify({"error": "export not found"}), 404
        if state.get("session_id") != session_id:
            return "", 403
        if state["status"] != "pending":
            return jsonify({"status": state["status"]})

        os.makedirs(_EXPORTS_DIR, exist_ok=True)
        lock_path = os.path.join(_EXPORTS_DIR, f"{export_id}.lock")

        # O_EXCL lock + stale 처리 (재시도 1회)
        for attempt in range(2):
            try:
                fd = os.open(lock_path, os.O_CREAT | os.O_EXCL | os.O_WRONLY)
                os.write(fd, str(os.getpid()).encode())
                os.close(fd)
                break
            except FileExistsError:
                if attempt == 0:
                    try:
                        if time.time() - os.path.getmtime(lock_path) > 300:
                            os.unlink(lock_path)
                            continue
                    except Exception:
                        pass
                return jsonify({"status": "running"})

        ev = threading.Event()
        _DOCX_TASKS[export_id] = ev
        threading.Thread(target=_run_docx_export, args=(export_id,), daemon=True).start()
        return jsonify({"status": "started"})
    except Exception as e:
        logger.error(f"❌ /docx-start 오류: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500


@app.route('/docx-stream/<export_id>')
def docx_stream(export_id):
    """DOCX export 진행 상태 SSE 스트림"""
    if not _is_valid_uuid(export_id):
        return "", 400

    user_name = request.args.get("user_name", "")
    session_id = f"{RAGHelperMixin.SESSION_ID_PREFIX}{user_name}"
    state = _read_export_state(export_id)
    if not state:
        return "", 404
    if state.get("session_id") != session_id:
        return "", 403

    STATUS_MESSAGES = {
        "pending":          (1, "[DOCX] 생성 준비 중..."),
        "running":          (1, "[DOCX] 변환 중..."),
        "generating_docx":  (2, "[DOCX] 파일 저장 중..."),
        "complete":         (2, "[DOCX] 생성 완료"),
        "error":            (2, "[DOCX] 생성 실패"),
    }

    def generate():
        last_status = None
        heartbeat_t = time.time()
        while True:
            try:
                s = _read_export_state(export_id)
            except Exception:
                yield f"event: error\ndata: {_json.dumps({'message': 'state read error'})}\n\n"
                return

            status = s.get("status", "pending")
            if status != last_status:
                step, msg = STATUS_MESSAGES.get(status, (1, status))
                payload = _json.dumps({"step": step, "total": 3, "message": msg}, ensure_ascii=False)
                yield f"event: progress\ndata: {payload}\n\n"
                last_status = status

            if status == "complete":
                url = f"/download-export/{export_id}"
                yield f"event: complete\ndata: {_json.dumps({'docx_url': url})}\n\n"
                return
            if status == "error":
                yield f"event: error\ndata: {_json.dumps({'message': s.get('error', '알 수 없는 오류')})}\n\n"
                return

            # heartbeat (30s 마다)
            if time.time() - heartbeat_t > 30:
                yield ": heartbeat\n\n"
                heartbeat_t = time.time()

            ev = _DOCX_TASKS.get(export_id)
            if ev:
                ev.wait(timeout=1.0)
                ev.clear()
            else:
                time.sleep(0.5)

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route('/download-export/<export_id>')
def download_export(export_id):
    """완료된 DOCX 파일 다운로드"""
    if not _is_valid_uuid(export_id):
        return "", 400

    user_name = request.args.get("user_name", "")
    session_id = f"{RAGHelperMixin.SESSION_ID_PREFIX}{user_name}"
    state = _read_export_state(export_id)
    if not state:
        return "", 404
    if state.get("session_id") != session_id:
        return "", 403

    conv_id = state.get("conversation_id")
    if conv_id and rag_engine.is_conversation_deleted(session_id, conv_id):
        return "", 403

    if state.get("status") != "complete":
        return "", 409

    real_exports = os.path.realpath(_EXPORTS_DIR)
    safe_path = os.path.realpath(os.path.join(_EXPORTS_DIR, f"{export_id}.docx"))
    if not safe_path.startswith(real_exports + os.sep):
        return "", 403
    if not os.path.exists(safe_path):
        return "", 404

    from flask import send_file as _send_file
    return _send_file(
        safe_path,
        as_attachment=True,
        download_name=f"export_{export_id}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route('/exports/pending')
def exports_pending():
    """새로고침 복원용: 현재 세션의 진행중 export 목록"""
    user_name = request.args.get("user_name", "")
    session_id = f"{RAGHelperMixin.SESSION_ID_PREFIX}{user_name}"

    os.makedirs(_EXPORTS_DIR, exist_ok=True)
    pending_list = []
    uuid_re = re.compile(r'^[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}\.json$', re.IGNORECASE)

    for fname in os.listdir(_EXPORTS_DIR):
        if not uuid_re.match(fname):
            continue
        try:
            with open(os.path.join(_EXPORTS_DIR, fname)) as f:
                s = _json.load(f)
        except Exception:
            continue
        if s.get("session_id") != session_id:
            continue
        if s.get("status") not in ("pending", "running", "generating_docx"):
            continue
        pending_list.append({
            "export_id": s.get("export_id"),
            "conversation_id": s.get("conversation_id"),
            "turn_index": s.get("turn_index"),
            "status": s.get("status"),
        })

    return jsonify(pending_list)


if __name__ == '__main__':
    logger.info("🌐 웹 서버 시작: http://0.0.0.0:5000")
    app.run(host='0.0.0.0', port=5000, debug=False)
